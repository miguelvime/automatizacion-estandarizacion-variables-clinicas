# -*- coding: utf-8 -*-
"""
===============================================================================
GENERADOR DE TABLAS, FIGURAS Y LISTADOS (TFL) - SECCIÓN 6.6.1 / 7.1 FIABILIDAD
===============================================================================
Genera los artefactos estadísticos listos para publicación médica / TFM,
con exportación directa a:
- Formatos tabulares: DOCX (Word con estilo APA/tablas limpias), XLSX, CSV, Markdown
- Formatos gráficos: PNG (300 DPI), SVG (vectorial), PDF
- Listados clínicos de discrepancias inter-iteraciones
"""

import json
from pathlib import Path
import matplotlib.pyplot as plt
import matplotlib as mpl
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Configuración de rutas
BASE_DIR = Path(__file__).resolve().parents[2]
RESULTS_DIR = BASE_DIR / "results"
LLM_DIR = RESULTS_DIR / "llm_text"
TFL_DIR = RESULTS_DIR / "TFL"
DIR_TABLAS = TFL_DIR / "tablas"
DIR_FIGURAS = TFL_DIR / "figuras"
DIR_LISTADOS = TFL_DIR / "listados"
DIR_INFORMES = TFL_DIR / "informes"

for d in [TFL_DIR, DIR_TABLAS, DIR_FIGURAS, DIR_LISTADOS]:
    d.mkdir(parents=True, exist_ok=True)

MODELOS = [
    {
        "id": "gemma_local",
        "nombre": "Gemma-4-31B-it (Local)",
        "tipo": "Local (On-Premise)",
        "archivo": LLM_DIR / "2026-08-11_gemma-4-31b-it-codified.json"
    },
    {
        "id": "gemini_35_cloud",
        "nombre": "Gemini Flash 3.5 (Cloud - Línea Base)",
        "tipo": "Cloud (API)",
        "archivo": LLM_DIR / "2026-08-18_gemini-flash-3.5_codified.json"
    },
    {
        "id": "gemini_36_cloud",
        "nombre": "Gemini Flash 3.6 (Cloud)",
        "tipo": "Cloud (API)",
        "archivo": LLM_DIR / "2026-08-18_gemini-flash-3.6_codified.json"
    }
]


def obtener_codigos_core_set(modelos):
    codigos = set()
    for m in modelos:
        if m["archivo"].exists():
            with open(m["archivo"], "r", encoding="utf-8") as f:
                datos = json.load(f)
                for item in datos:
                    codigos.update(item.get("icf_codes", []))
                    codigos.update(item.get("predicted_icf_it1", []))
                    codigos.update(item.get("predicted_icf_it2", []))
                    codigos.update(item.get("predicted_icf_it3", []))
    return sorted(list(codigos))


def analizar_modelo(ruta_archivo: Path, nombre_modelo: str, tipo_modelo: str, codigos_cif: list):
    with open(ruta_archivo, "r", encoding="utf-8") as f:
        historias = json.load(f)

    total_historias = len(historias)
    acuerdos_exactos_historia = 0
    discrepancias = []

    for i, historia in enumerate(historias, start=1):
        it1 = set(historia.get("predicted_icf_it1", []))
        it2 = set(historia.get("predicted_icf_it2", []))
        it3 = set(historia.get("predicted_icf_it3", []))

        if it1 == it2 == it3:
            acuerdos_exactos_historia += 1
        else:
            discrepancias.append({
                "historia_idx": i,
                "historia_id": historia.get("id", f"H_{i:03d}"),
                "it1": sorted(list(it1)),
                "it2": sorted(list(it2)),
                "it3": sorted(list(it3)),
                "union": sorted(list(it1 | it2 | it3)),
                "interseccion": sorted(list(it1 & it2 & it3))
            })

    emr_pct = (acuerdos_exactos_historia / total_historias) * 100.0

    matriz_votos = []
    filas_111 = 0
    filas_000 = 0
    filas_desacuerdo = 0

    for historia in historias:
        it1 = set(historia.get("predicted_icf_it1", []))
        it2 = set(historia.get("predicted_icf_it2", []))
        it3 = set(historia.get("predicted_icf_it3", []))

        for codigo in codigos_cif:
            v1 = 1 if codigo in it1 else 0
            v2 = 1 if codigo in it2 else 0
            v3 = 1 if codigo in it3 else 0
            matriz_votos.append([v1, v2, v3])

            suma = v1 + v2 + v3
            if suma == 3:
                filas_111 += 1
            elif suma == 0:
                filas_000 += 1
            else:
                filas_desacuerdo += 1

    N = len(matriz_votos)  # 3078
    K = 3

    sumatoria_acuerdo = 0.0
    total_unos = 0
    total_ceros = 0

    for fila in matriz_votos:
        n1 = sum(fila)
        n0 = K - n1
        total_unos += n1
        total_ceros += n0
        acuerdo_fila = (n1 * (n1 - 1) + n0 * (n0 - 1)) / (K * (K - 1))
        sumatoria_acuerdo += acuerdo_fila

    Po = sumatoria_acuerdo / N

    p1 = total_unos / (N * K)
    p0 = total_ceros / (N * K)
    Pe_gwet = 2.0 * p1 * p0
    gwet_ac1 = (Po - Pe_gwet) / (1.0 - Pe_gwet) if (1.0 - Pe_gwet) != 0 else 1.0

    Do = 1.0 - Po
    T = N * K
    De = (2.0 * total_ceros * total_unos) / (T * (T - 1)) if T > 1 else 0.0
    kripp_alpha = 1.0 - (Do / De) if De != 0 else 1.0

    return {
        "nombre": nombre_modelo,
        "tipo": tipo_modelo,
        "historias": total_historias,
        "emr_acuerdos": acuerdos_exactos_historia,
        "emr_pct": emr_pct,
        "unidades_totales": N,
        "filas_111": filas_111,
        "filas_000": filas_000,
        "filas_desacuerdo": filas_desacuerdo,
        "pct_111": (filas_111 / N) * 100.0,
        "pct_000": (filas_000 / N) * 100.0,
        "pct_desacuerdo": (filas_desacuerdo / N) * 100.0,
        "Po": Po,
        "Po_pct": Po * 100.0,
        "Gwet_AC1": gwet_ac1,
        "Krippendorff_Alpha": kripp_alpha,
        "discrepancias": discrepancias
    }


def estilizar_tabla_word(table, col_widths=None):
    """Aplica formato APA / Medical Journal estándar a una tabla Word."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Estilo de bordes APA: línea superior e inferior en encabezado, línea inferior en tabla
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="1A365D"/>\n'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="1A365D"/>\n'
        f'  <w:insideH w:val="none"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Cabecera
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
        bottom_border = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="0" w:color="1A365D"/></w:tcBorders>')
        tcPr.append(shd)
        tcPr.append(bottom_border)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(26, 54, 93)

    # Filas de datos
    for row_idx, row in enumerate(table.rows[1:], start=1):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Opcional zebra striping muy sutil
            if row_idx % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                tcPr.append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)


def generar_figuras(resultados):
    """Genera figuras en alta resolución (PNG 300 DPI, SVG y PDF)."""
    # Estilo general
    plt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')
    mpl.rcParams['font.family'] = 'DejaVu Sans'
    mpl.rcParams['font.size'] = 10
    mpl.rcParams['axes.titlesize'] = 12
    mpl.rcParams['axes.labelsize'] = 10

    # -------------------------------------------------------------
    # FIGURA 1: Comparativa de Métricas de Confiabilidad Inter-Iteraciones
    # -------------------------------------------------------------
    fig, ax = plt.subplots(figsize=(8.5, 4.8), dpi=300)
    
    nombres_cortos = ["Gemma-4-31B (Local)", "Gemini Flash 3.5 (Cloud)", "Gemini Flash 3.6 (Cloud)"]
    x = np.arange(len(nombres_cortos))
    width = 0.20

    emr_vals = [r["emr_pct"] for r in resultados]
    po_vals = [r["Po_pct"] for r in resultados]
    gwet_vals = [r["Gwet_AC1"] * 100 for r in resultados]
    kripp_vals = [r["Krippendorff_Alpha"] * 100 for r in resultados]

    colors = ['#1E3A8A', '#0D9488', '#F59E0B', '#10B981']

    rects1 = ax.bar(x - 1.5*width, emr_vals, width, label='Exact Match Paciente (%)', color=colors[0], edgecolor='white', linewidth=1)
    rects2 = ax.bar(x - 0.5*width, po_vals, width, label='Acuerdo Observado Po (%)', color=colors[1], edgecolor='white', linewidth=1)
    rects3 = ax.bar(x + 0.5*width, gwet_vals, width, label="Gwet's AC1 (x100)", color=colors[2], edgecolor='white', linewidth=1)
    rects4 = ax.bar(x + 1.5*width, kripp_vals, width, label='Krippendorff α (x100)', color=colors[3], edgecolor='white', linewidth=1)

    ax.set_ylabel('Porcentaje / Índice escalado (%)', fontweight='bold')
#     ax.set_title('Confiabilidad y Reproducibilidad Inter-Iteraciones (K = 3) por Modelo LLM', fontsize=12, pad=15, fontweight='bold', color='#1E293B')
    ax.set_xticks(x)
    ax.set_xticklabels(nombres_cortos, fontweight='bold')
    ax.set_ylim(94, 101.5)
    ax.axhline(100, color='gray', linestyle='--', linewidth=0.8, alpha=0.7)
    ax.legend(loc='lower right', frameon=True, facecolor='white', framealpha=0.95, fontsize=9)

    # Etiquetas sobre las barras
    def autolabel(rects, is_alpha=False):
        for rect in rects:
            height = rect.get_height()
            txt = f"{height:.2f}%" if not is_alpha else f"{height/100:.4f}"
            ax.annotate(f"{height:.2f}",
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  
                        textcoords="offset points",
                        ha='center', va='bottom', fontsize=7.5, rotation=0, fontweight='bold')

    autolabel(rects1)
    autolabel(rects2)
    autolabel(rects3)
    autolabel(rects4)

    plt.tight_layout()
    for ext in ['png', 'svg', 'pdf']:
        fig.savefig(DIR_FIGURAS / f"figura1_comparativa_fiabilidad_modelos.{ext}", dpi=300, bbox_inches='tight')
    plt.close(fig)

    # -------------------------------------------------------------
    # FIGURA 2: Desglose del Espacio de Decisiones Binarias (3.078 U.)
    # -------------------------------------------------------------
    fig, ax = plt.subplots(figsize=(8.0, 4.2), dpi=300)
    
    n_modelos = len(resultados)
    y_pos = np.arange(n_modelos)
    
    pct_111 = [r["pct_111"] for r in resultados]
    pct_000 = [r["pct_000"] for r in resultados]
    pct_disc = [r["pct_desacuerdo"] for r in resultados]

    bar_h = 0.45
    b1 = ax.barh(y_pos, pct_000, bar_h, label='Abstención Unánime [0,0,0] (Ausencia)', color='#3B82F6', alpha=0.9)
    b2 = ax.barh(y_pos, pct_111, bar_h, left=pct_000, label='Asignación Unánime [1,1,1] (Presencia)', color='#10B981', alpha=0.9)
    b3 = ax.barh(y_pos, pct_disc, bar_h, left=np.array(pct_000)+np.array(pct_111), label='Discrepancia Inter-pasada (<3/3)', color='#EF4444')

    ax.set_yticks(y_pos)
    ax.set_yticklabels(nombres_cortos, fontweight='bold')
    ax.set_xlabel('Distribución Porcentual en Matriz Ontológica Completa (%)', fontweight='bold')
#     ax.set_title('Composición del Espacio de Decisiones Binarias (114 historias × 27 códigos CIF = 3.078)', fontsize=11, fontweight='bold', pad=12)
    ax.set_xlim(0, 100.5)
    ax.legend(loc='lower center', bbox_to_anchor=(0.5, -0.28), ncol=3, frameon=True, fontsize=8.5)

    # Anotar porcentajes
    for i in range(n_modelos):
        ax.text(pct_000[i]/2, i, f"{pct_000[i]:.1f}%", ha='center', va='center', color='white', fontweight='bold', fontsize=8.5)
        ax.text(pct_000[i] + pct_111[i]/2, i, f"{pct_111[i]:.1f}%", ha='center', va='center', color='white', fontweight='bold', fontsize=8.5)
        if pct_disc[i] > 0:
            ax.text(99.5, i, f"Disc: {pct_disc[i]:.2f}%", ha='right', va='center', color='#991B1B', fontweight='bold', fontsize=7.5)
        else:
            ax.text(99.5, i, "0% Disc.", ha='right', va='center', color='#065F46', fontweight='bold', fontsize=7.5)

    plt.tight_layout()
    for ext in ['png', 'svg', 'pdf']:
        fig.savefig(DIR_FIGURAS / f"figura2_espacio_decisiones_binarias.{ext}", dpi=300, bbox_inches='tight')
    plt.close(fig)

    print("✅ Figuras generadas correctamente en formatos PNG (300 DPI), SVG y PDF.")


def generar_tablas_y_documentos(resultados, codigos_cif):
    """Crea DataFrames, CSV, XLSX y documentos DOCX listos para copiar/pegar en Word."""
    
    # -------------------------------------------------------------
    # TABLA 1: Resumen de Confiabilidad y Reproducibilidad
    # -------------------------------------------------------------
    datos_t1 = []
    for r in resultados:
        datos_t1.append({
            "Modelo LLM": r["nombre"],
            "Despliegue": r["tipo"],
            "Historias Evaluadas": r["historias"],
            "Acuerdo Exacto 3/3 (n)": f"{r['emr_acuerdos']}/{r['historias']}",
            "Exact Match (%)": f"{r['emr_pct']:.2f}%",
            "Acuerdo Observado Po (%)": f"{r['Po_pct']:.4f}%",
            "Gwet's AC1": f"{r['Gwet_AC1']:.4f}",
            "Krippendorff α": f"{r['Krippendorff_Alpha']:.4f}"
        })
    df_t1 = pd.DataFrame(datos_t1)
    df_t1.to_csv(DIR_TABLAS / "tabla1_fiabilidad_inter_iteraciones.csv", index=False, encoding='utf-8-sig')
    df_t1.to_excel(DIR_TABLAS / "tabla1_fiabilidad_inter_iteraciones.xlsx", index=False)

    # -------------------------------------------------------------
    # TABLA 2: Matriz Ontológica y Frecuencias de Decisiones Binarias
    # -------------------------------------------------------------
    datos_t2 = []
    for r in resultados:
        datos_t2.append({
            "Modelo LLM": r["nombre"],
            "Total Decisiones Binarias": r["unidades_totales"],
            "Asignación Unánime SÍ [1,1,1]": f"{r['filas_111']} ({r['pct_111']:.2f}%)",
            "Abstención Unánime NO [0,0,0]": f"{r['filas_000']} ({r['pct_000']:.2f}%)",
            "Discrepancia en Iteración": f"{r['filas_desacuerdo']} ({r['pct_desacuerdo']:.2f}%)",
            "Determinismo Paciente": f"{r['emr_acuerdos']}/{r['historias']} ({r['emr_pct']:.2f}%)"
        })
    df_t2 = pd.DataFrame(datos_t2)
    df_t2.to_csv(DIR_TABLAS / "tabla2_matriz_decisiones_binarias.csv", index=False, encoding='utf-8-sig')
    df_t2.to_excel(DIR_TABLAS / "tabla2_matriz_decisiones_binarias.xlsx", index=False)

    # -------------------------------------------------------------
    # LISTADO 1: Detalle Clínico de Discrepancias
    # -------------------------------------------------------------
    filas_listing = []
    for r in resultados:
        if r["discrepancias"]:
            for d in r["discrepancias"]:
                filas_listing.append({
                    "Modelo": r["nombre"],
                    "Historia Clínica": f"Historia #{d['historia_idx']}",
                    "Iteración 1": ", ".join(d["it1"]) if d["it1"] else "(Vacío)",
                    "Iteración 2": ", ".join(d["it2"]) if d["it2"] else "(Vacío)",
                    "Iteración 3": ", ".join(d["it3"]) if d["it3"] else "(Vacío)",
                    "Códigos en Consenso (3/3)": ", ".join(d["interseccion"]) if d["interseccion"] else "(Ninguno)",
                    "Código(s) Discrepante(s)": ", ".join(sorted(list(set(d["union"]) - set(d["interseccion"]))))
                })
        else:
            filas_listing.append({
                "Modelo": r["nombre"],
                "Historia Clínica": "Todas (114/114)",
                "Iteración 1": "Acuerdo perfecto",
                "Iteración 2": "Acuerdo perfecto",
                "Iteración 3": "Acuerdo perfecto",
                "Códigos en Consenso (3/3)": "100% Determinista",
                "Código(s) Discrepante(s)": "Ninguno (0 discrepancias)"
            })
    df_listing = pd.DataFrame(filas_listing)
    df_listing.to_csv(DIR_LISTADOS / "listing1_discrepancias_detalladas.csv", index=False, encoding='utf-8-sig')
    df_listing.to_excel(DIR_LISTADOS / "listing1_discrepancias_detalladas.xlsx", index=False)

    # -------------------------------------------------------------
    # GENERACIÓN DE DOCUMENTO WORD (DOCX) PROFESIONAL
    # -------------------------------------------------------------
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Título principal
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Tablas y Resultados de Fiabilidad Inter-Iteraciones (Sección 6.6.1 / 7.1)")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(26, 54, 93)
    p_title.paragraph_format.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("TFM: Codificación Automatizada de Texto Clínico No Estructurado a Estándares CIF mediante Procesamiento de Lenguaje Natural")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(10)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)
    p_sub.paragraph_format.space_after = Pt(18)

    # Texto introductorio
    p_desc = doc.add_paragraph()
    p_desc.add_run(
        "A continuación se presentan las tablas estadísticas con formato estándar para publicación académica / inclusión directa en Microsoft Word. "
        "Las tablas han sido calculadas sobre el corpus de 114 historias clínicas y el espacio ontológico de 27 códigos CIF del Core Set de Dolor Crónico Generalizado (3.078 decisiones binarias por modelo, K = 3 iteraciones independientes)."
    )
    p_desc.paragraph_format.space_after = Pt(14)

    # ------------------ TABLA 1 EN WORD ------------------
    p_t1_title = doc.add_paragraph()
    r_t1_tag = p_t1_title.add_run("Tabla 1. ")
    r_t1_tag.bold = True
    r_t1_title = p_t1_title.add_run("Confiabilidad y reproducibilidad inter-iteraciones (K = 3) en la codificación CIF.")
    p_t1_title.paragraph_format.space_after = Pt(6)

    t1_word = doc.add_table(rows=len(datos_t1)+1, cols=len(df_t1.columns))
    for col_idx, col_name in enumerate(df_t1.columns):
        t1_word.cell(0, col_idx).text = col_name

    for row_idx, row_data in enumerate(datos_t1, start=1):
        for col_idx, col_name in enumerate(df_t1.columns):
            t1_word.cell(row_idx, col_idx).text = str(row_data[col_name])

    estilizar_tabla_word(t1_word, col_widths=[2.1, 1.1, 0.8, 0.9, 0.8, 0.9, 0.7, 0.7])

    p_t1_note = doc.add_paragraph()
    p_t1_note.paragraph_format.space_before = Pt(4)
    p_t1_note.paragraph_format.space_after = Pt(16)
    r_note1 = p_t1_note.add_run("Nota: EMR = Exact Match Ratio a nivel paciente (conjunto idéntico de códigos en las 3 iteraciones); Po = Porcentaje de acuerdo observado sobre 3.078 decisiones binarias; Gwet's AC1 = Coeficiente de concordancia inter-evaluador corregido por azar y robusto al desbalanceo; Krippendorff α = Coeficiente nominal de reproducibilidad multievaluador.")
    r_note1.font.size = Pt(8)
    r_note1.font.italic = True
    r_note1.font.color.rgb = RGBColor(100, 116, 139)

    # ------------------ TABLA 2 EN WORD ------------------
    p_t2_title = doc.add_paragraph()
    r_t2_tag = p_t2_title.add_run("Tabla 2. ")
    r_t2_tag.bold = True
    r_t2_title = p_t2_title.add_run("Desglose del espacio de decisiones binarias ontológicas (114 historias × 27 códigos = 3.078 unidades).")
    p_t2_title.paragraph_format.space_after = Pt(6)

    t2_word = doc.add_table(rows=len(datos_t2)+1, cols=len(df_t2.columns))
    for col_idx, col_name in enumerate(df_t2.columns):
        t2_word.cell(0, col_idx).text = col_name

    for row_idx, row_data in enumerate(datos_t2, start=1):
        for col_idx, col_name in enumerate(df_t2.columns):
            t2_word.cell(row_idx, col_idx).text = str(row_data[col_name])

    estilizar_tabla_word(t2_word, col_widths=[2.0, 0.9, 1.2, 1.2, 1.1, 1.1])

    p_t2_note = doc.add_paragraph()
    p_t2_note.paragraph_format.space_before = Pt(4)
    p_t2_note.paragraph_format.space_after = Pt(16)
    r_note2 = p_t2_note.add_run("Nota: Las asignaciones unánimes en SÍ [1,1,1] reflejan codificación afirmativa consistente; las abstenciones unánimes en NO [0,0,0] reflejan abstención diagnóstica correcta en códigos no presentes en la narrativa. Las discrepancias corresponden a códigos asignados en 1 o 2 de las 3 pasadas.")
    r_note2.font.size = Pt(8)
    r_note2.font.italic = True
    r_note2.font.color.rgb = RGBColor(100, 116, 139)

    # ------------------ LISTADO 1 EN WORD ------------------
    p_l1_title = doc.add_paragraph()
    r_l1_tag = p_l1_title.add_run("Listado 1. ")
    r_l1_tag.bold = True
    r_l1_title = p_l1_title.add_run("Auditoría clínica de historias con discrepancias inter-iteraciones (K = 3).")
    p_l1_title.paragraph_format.space_after = Pt(6)

    t3_word = doc.add_table(rows=len(filas_listing)+1, cols=len(df_listing.columns))
    for col_idx, col_name in enumerate(df_listing.columns):
        t3_word.cell(0, col_idx).text = col_name

    for row_idx, row_data in enumerate(filas_listing, start=1):
        for col_idx, col_name in enumerate(df_listing.columns):
            t3_word.cell(row_idx, col_idx).text = str(row_data[col_name])

    estilizar_tabla_word(t3_word, col_widths=[1.5, 0.9, 1.1, 1.1, 1.1, 1.1, 1.2])

    p_l1_note = doc.add_paragraph()
    p_l1_note.paragraph_format.space_before = Pt(4)
    p_l1_note.paragraph_format.space_after = Pt(16)
    r_note3 = p_l1_note.add_run("Nota: La estrategia de autoconsistencia (3/3 o voto mayoritario 2/3) resuelve automáticamente estas variaciones puntuales asegurando la reproducibilidad operativa.")
    r_note3.font.size = Pt(8)
    r_note3.font.italic = True
    r_note3.font.color.rgb = RGBColor(100, 116, 139)

    doc_path = DIR_TABLAS / "tablas_completas_fiabilidad_word.docx"
    doc.save(doc_path)
    print(f"✅ Documento Word con tablas APA guardado en: {doc_path}")


def main():
    codigos_cif = obtener_codigos_core_set(MODELOS)
    print(f"📊 Espacio Ontológico Core Set CIF: {len(codigos_cif)} códigos únicos.")
    
    resultados = []
    for m in MODELOS:
        if m["archivo"].exists():
            res = analizar_modelo(m["archivo"], m["nombre"], m["tipo"], codigos_cif)
            resultados.append(res)

    generar_figuras(resultados)
    generar_tablas_y_documentos(resultados, codigos_cif)

    # Generar Informe Markdown Consolidado
    informe_md_path = DIR_INFORMES / "INFORME_TFL_FIABILIDAD.md"
    with open(informe_md_path, "w", encoding="utf-8") as f:
        f.write("# Informe de Tablas, Figuras y Listados (TFL): Fiabilidad Inter-Iteraciones (Sección 6.6.1 / 7.1)\n\n")
        f.write("Este informe contiene los resultados cuantitativos definitivos, tablas formateadas para publicación médica y gráficos vectoriales/alta resolución correspondientes a la **Sección 6.6.1 (Metodología)** y **Sección 7.1 (Resultados)** del TFM.\n\n")
        
        f.write("## 1. Tabla 1: Confiabilidad y Reproducibilidad Inter-Iteraciones (K = 3)\n\n")
        f.write("| Modelo LLM Evaluado | Tipo de Despliegue | Historias | Exact Match (3/3) | Acuerdo Po (%) | Gwet's AC1 | Krippendorff α |\n")
        f.write("| :--- | :--- | :---: | :---: | :---: | :---: | :---: |\n")
        for r in resultados:
            f.write(f"| **{r['nombre']}** | {r['tipo']} | {r['historias']} | {r['emr_acuerdos']}/{r['historias']} ({r['emr_pct']:.2f}%) | {r['Po_pct']:.4f}% | {r['Gwet_AC1']:.4f} | {r['Krippendorff_Alpha']:.4f} |\n")
        f.write("\n> **Interpretación metodológica:** Un valor de $\\alpha > 0.80$ y $AC1 > 0.80$ denota acuerdo casi perfecto (Landis & Koch / Krippendorff). Tanto el despliegue local (Gemma) como en la nube (Gemini Flash) demuestran determinismo operativo virtualmente perfecto ($>0.998$).\n\n")

        f.write("## 2. Tabla 2: Desglose del Espacio de Decisiones Binarias (3.078 Unidades Ontológicas)\n\n")
        f.write("| Modelo LLM | Unidades Totales | Asignación Unánime SÍ [1,1,1] | Abstención Unánime NO [0,0,0] | Discrepancias (<3/3) | Determinismo Paciente |\n")
        f.write("| :--- | :---: | :---: | :---: | :---: | :---: |\n")
        for r in resultados:
            f.write(f"| **{r['nombre']}** | {r['unidades_totales']} | {r['filas_111']} ({r['pct_111']:.2f}%) | {r['filas_000']} ({r['pct_000']:.2f}%) | {r['filas_desacuerdo']} ({r['pct_desacuerdo']:.2f}%) | {r['emr_acuerdos']}/{r['historias']} ({r['emr_pct']:.2f}%) |\n")
        f.write("\n\n")

        f.write("## 3. Listado 1: Auditoría de Historias Clínicas con Discrepancias\n\n")
        f.write("| Modelo | Historia | Iteración 1 | Iteración 2 | Iteración 3 | Consenso (3/3) | Código Discrepante |\n")
        f.write("| :--- | :---: | :--- | :--- | :--- | :--- | :--- |\n")
        for r in resultados:
            if r["discrepancias"]:
                for d in r["discrepancias"]:
                    diff = sorted(list(set(d["union"]) - set(d["interseccion"])))
                    f.write(f"| {r['nombre']} | #{d['historia_idx']} | {', '.join(d['it1'])} | {', '.join(d['it2'])} | {', '.join(d['it3'])} | {', '.join(d['interseccion'])} | `{', '.join(diff)}` |\n")
            else:
                f.write(f"| {r['nombre']} | Todas (114/114) | - | - | - | 100% Determinista | Ninguno (0) |\n")
        f.write("\n\n")

        f.write("## 4. Figuras Generadas para Publicación\n\n")
        f.write("- **Figura 1**: `results/stats_v4/TFL/figuras/figura1_comparativa_fiabilidad_modelos.png` (PNG 300 DPI, SVG, PDF)\n")
        f.write("- **Figura 2**: `results/stats_v4/TFL/figuras/figura2_espacio_decisiones_binarias.png` (PNG 300 DPI, SVG, PDF)\n\n")

    print(f"✅ Informe Markdown guardado en: {informe_md_path}")


if __name__ == "__main__":
    main()
