# -*- coding: utf-8 -*-
"""
===============================================================================
CÁLCULO DE DESEMPEÑO DIAGNÓSTICO: F1-SCORE (MICRO, MACRO, WEIGHTED Y PER CLASS)
EVALUACIÓN COMPARATIVA FRENTE AL GROUND TRUTH (CORE SET CIF - 27 CATEGORÍAS)
===============================================================================

¿Qué calcula este script?
-------------------------
Evalúa la validez diagnóstica de los modelos LLM (Gemma-4-31B-it, Gemini Flash 3.5,
Gemini Flash 3.6) comparando las predicciones consolidadas bajo consenso estricto (3/3)
frente al Ground Truth (códigos generadores `icf_codes`) en las 114 historias clínicas.

Estructura de Presentación Científica (Estilo Publicación Médica):
- Tabla Principal Transpuesta: Modelos en Columnas y Variables/Métricas en Filas
  (Parámetros de Corpus, Matriz de Confusión, Exact Match, Micro, Macro y Weighted F1 con IC 95%).
- Tabla de Auditoría Detallada: Las 27 Categorías CIF del Core Set de Dolor Crónico.

Salidas generadas en `results/stats_v4/`:
- Documento Word (.docx): `tablas_desempeno.docx` con tablas nativas editables y diseño APA.
- Archivo JSON: `resumen_f1_score.json` con todos los estadísticos descriptivos.
"""

import json
import random
from pathlib import Path
from typing import Dict, List, Any, Set, Tuple

# Importaciones para la exportación nativa a Word (.docx)
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls
    TIENE_DOCX = True
except ImportError:
    TIENE_DOCX = False

# Rutas de datos
BASE_DIR = Path(__file__).resolve().parents[2]
LLM_DIR = BASE_DIR / "results" / "llm_text"
TFL_DIR = BASE_DIR / "results" / "TFL"
TABLAS_DIR = TFL_DIR / "tablas"

MODELOS = [
    {
        "id": "gemma_31b",
        "nombre": "Gemma-4-31B-it",
        "archivo": LLM_DIR / "2026-08-11_gemma-4-31b-it-codified.json"
    },
    {
        "id": "gemini_flash_35",
        "nombre": "Gemini Flash 3.5",
        "archivo": LLM_DIR / "2026-08-18_gemini-flash-3.5_codified.json"
    },
    {
        "id": "gemini_flash_36",
        "nombre": "Gemini Flash 3.6",
        "archivo": LLM_DIR / "2026-08-18_gemini-flash-3.6_codified.json"
    }
]

# Diccionario oficial de las 27 categorías CIF del Core Set de Dolor Crónico Generalizado
DICCIONARIO_CIF = {
    # Funciones Corporales (b)
    "b130": "Funciones relacionadas con la energía y los impulsos",
    "b134": "Funciones del sueño",
    "b147": "Funciones psicomotoras",
    "b152": "Funciones emocionales",
    "b1602": "Contenido del pensamiento",
    "b175": "Funciones cognitivas superiores (Resolver problemas)",
    "b240": "Sensaciones corporales y manejo del estrés",
    "b280": "Sensación de dolor",
    "b455": "Tolerancia al ejercicio físico",
    "b730": "Funciones relacionadas con la fuerza muscular",
    "b760": "Control de los movimientos voluntarios",
    # Actividades y Participación (d)
    "d175": "Resolver problemas",
    "d230": "Llevar a cabo rutinas diarias",
    "d240": "Manejo del estrés y demandas psicológicas",
    "d290": "Tareas y demandas generales (Ocio)",
    "d430": "Levantar y llevar objetos",
    "d450": "Andar y desplazarse",
    "d640": "Realizar los quehaceres de la casa",
    "d760": "Relaciones familiares",
    "d770": "Relaciones íntimas y sociales",
    "d850": "Trabajo remunerado",
    "d920": "Tiempo libre y ocio",
    # Factores Ambientales (e)
    "e1101": "Medicamentos",
    "e310": "Familiares cercanos",
    "e355": "Profesionales de la salud",
    "e410": "Actitudes individuales de miembros de la familia cercana",
    "e570": "Servicios, sistemas y políticas de seguridad social"
}


def obtener_todos_los_codigos(modelos: List[Dict[str, Any]]) -> List[str]:
    """Obtiene la lista única ordenada de las 27 categorías CIF del Core Set."""
    codigos = set(DICCIONARIO_CIF.keys())
    for m in modelos:
        if m["archivo"].exists():
            with open(m["archivo"], "r", encoding="utf-8") as f:
                datos = json.load(f)
                for item in datos:
                    codigos.update(item.get("icf_codes", []))
                    codigos.update(item.get("predicted_icf_codes_consensus", []))
    return sorted(list(codigos))


def calcular_desempeno(historias: List[Dict[str, Any]], codigos_cif: List[str]) -> Dict[str, Any]:
    """
    Calcula la matriz de confusión y las métricas F1 (Micro, Macro, Weighted y Per Class).
    """
    n_historias = len(historias)
    total_tp = 0
    total_fp = 0
    total_fn = 0
    exact_matches = 0

    per_class = {
        c: {"tp": 0, "fp": 0, "fn": 0, "soporte": 0}
        for c in codigos_cif
    }

    for item in historias:
        gt = set(item.get("icf_codes", []))
        pred = set(item.get("predicted_icf_codes_consensus", []))

        if gt == pred:
            exact_matches += 1

        tp = len(gt & pred)
        fp = len(pred - gt)
        fn = len(gt - pred)

        total_tp += tp
        total_fp += fp
        total_fn += fn

        for c in codigos_cif:
            in_gt = c in gt
            in_pred = c in pred

            if in_gt and in_pred:
                per_class[c]["tp"] += 1
            elif in_pred and not in_gt:
                per_class[c]["fp"] += 1
            elif in_gt and not in_pred:
                per_class[c]["fn"] += 1

            if in_gt:
                per_class[c]["soporte"] += 1

    micro_p = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0.0
    micro_r = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0.0
    micro_f1 = (2 * micro_p * micro_r / (micro_p + micro_r)) if (micro_p + micro_r) > 0 else 0.0

    macro_p_sum = 0.0
    macro_r_sum = 0.0
    macro_f1_sum = 0.0

    weighted_p_sum = 0.0
    weighted_r_sum = 0.0
    weighted_f1_sum = 0.0

    soporte_total = sum(v["soporte"] for v in per_class.values())

    for c in codigos_cif:
        tp_c = per_class[c]["tp"]
        fp_c = per_class[c]["fp"]
        fn_c = per_class[c]["fn"]
        sup_c = per_class[c]["soporte"]

        prec_c = tp_c / (tp_c + fp_c) if (tp_c + fp_c) > 0 else 0.0
        rec_c = tp_c / (tp_c + fn_c) if (tp_c + fn_c) > 0 else 0.0
        f1_c = (2 * prec_c * rec_c / (prec_c + rec_c)) if (prec_c + rec_c) > 0 else 0.0

        per_class[c]["precision"] = prec_c
        per_class[c]["recall"] = rec_c
        per_class[c]["f1"] = f1_c
        per_class[c]["nombre"] = DICCIONARIO_CIF.get(c, c)

        macro_p_sum += prec_c
        macro_r_sum += rec_c
        macro_f1_sum += f1_c

        weighted_p_sum += prec_c * sup_c
        weighted_r_sum += rec_c * sup_c
        weighted_f1_sum += f1_c * sup_c

    n_clases = len(codigos_cif)
    macro_p = macro_p_sum / n_clases if n_clases > 0 else 0.0
    macro_r = macro_r_sum / n_clases if n_clases > 0 else 0.0
    macro_f1 = macro_f1_sum / n_clases if n_clases > 0 else 0.0

    weighted_p = weighted_p_sum / soporte_total if soporte_total > 0 else 0.0
    weighted_r = weighted_r_sum / soporte_total if soporte_total > 0 else 0.0
    weighted_f1 = weighted_f1_sum / soporte_total if soporte_total > 0 else 0.0

    emr_pct = (exact_matches / n_historias) * 100.0 if n_historias > 0 else 0.0

    return {
        "n_historias": n_historias,
        "n_clases": n_clases,
        "soporte_total": soporte_total,
        "exact_matches_n": exact_matches,
        "emr_pct": emr_pct,
        "confusion_global": {
            "tp": total_tp,
            "fp": total_fp,
            "fn": total_fn
        },
        "micro": {
            "precision": micro_p,
            "recall": micro_r,
            "f1": micro_f1
        },
        "macro": {
            "precision": macro_p,
            "recall": macro_r,
            "f1": macro_f1
        },
        "weighted": {
            "precision": weighted_p,
            "recall": weighted_r,
            "f1": weighted_f1
        },
        "por_clase": per_class
    }


def calcular_bootstrap_ci(historias: List[Dict[str, Any]], codigos_cif: List[str], n_iter: int = 1000, semilla: int = 2026) -> Dict[str, Tuple[float, float]]:
    random.seed(semilla)
    n_docs = len(historias)

    boot_micro_f1: List[float] = []
    boot_macro_f1: List[float] = []
    boot_weighted_f1: List[float] = []

    for _ in range(n_iter):
        muestra = [random.choice(historias) for _ in range(n_docs)]
        res = calcular_desempeno(muestra, codigos_cif)
        boot_micro_f1.append(res["micro"]["f1"])
        boot_macro_f1.append(res["macro"]["f1"])
        boot_weighted_f1.append(res["weighted"]["f1"])

    boot_micro_f1.sort()
    boot_macro_f1.sort()
    boot_weighted_f1.sort()

    idx_low = int(0.025 * n_iter)
    idx_high = int(0.975 * n_iter)

    return {
        "micro_f1": (boot_micro_f1[idx_low], boot_micro_f1[idx_high]),
        "macro_f1": (boot_macro_f1[idx_low], boot_macro_f1[idx_high]),
        "weighted_f1": (boot_weighted_f1[idx_low], boot_weighted_f1[idx_high])
    }


def construir_filas_tabla_transpuesta(resultados: List[Dict[str, Any]]) -> List[Tuple[str, str, str, str, bool]]:
    gemma = next(r for r in resultados if r["meta"]["id"] == "gemma_31b")
    g35 = next(r for r in resultados if r["meta"]["id"] == "gemini_flash_35")
    g36 = next(r for r in resultados if r["meta"]["id"] == "gemini_flash_36")

    filas = [
        ("PARÁMETROS DEL CORPUS CLÍNICO", "", "", "", True),
        ("Historias clínicas evaluadas (N)", "114", "114", "114", False),
        ("Espacio de decisiones ontológicas (114 × 27)", "3.078", "3.078", "3.078", False),
        ("Instancias CIF totales en Ground Truth (Soporte)", "465", "465", "465", False),
        ("Criterio de consenso inter-iteraciones", "Estricto (3/3)", "Estricto (3/3)", "Estricto (3/3)", False),
        
        ("MATRIZ DE CONFUSIÓN Y EXACTITUD", "", "", "", True),
        ("Verdaderos Positivos (TP)", f"{gemma['metricas']['confusion_global']['tp']}", f"{g35['metricas']['confusion_global']['tp']}", f"{g36['metricas']['confusion_global']['tp']}", False),
        ("Falsos Positivos (FP / Alucinaciones)", f"{gemma['metricas']['confusion_global']['fp']}", f"{g35['metricas']['confusion_global']['fp']}", f"{g36['metricas']['confusion_global']['fp']}", False),
        ("Falsos Negativos (FN / Omisiones)", f"{gemma['metricas']['confusion_global']['fn']}", f"{g35['metricas']['confusion_global']['fn']}", f"{g36['metricas']['confusion_global']['fn']}", False),
        ("Exact Match Ratio - EMR (n / N)", f"{gemma['metricas']['exact_matches_n']}/114", f"{g35['metricas']['exact_matches_n']}/114", f"{g36['metricas']['exact_matches_n']}/114", False),
        ("Exact Match Ratio - EMR (%)", f"{gemma['metricas']['emr_pct']:.2f}%", f"{g35['metricas']['emr_pct']:.2f}%", f"{g36['metricas']['emr_pct']:.2f}%", False),
        
        ("EFICACIA DIAGNÓSTICA GLOBAL (NIVEL MICRO)", "", "", "", True),
        ("Precisión Micro", f"{gemma['metricas']['micro']['precision']:.4f} ({gemma['metricas']['micro']['precision']*100:.2f}%)", f"{g35['metricas']['micro']['precision']:.4f} ({g35['metricas']['micro']['precision']*100:.2f}%)", f"{g36['metricas']['micro']['precision']:.4f} ({g36['metricas']['micro']['precision']*100:.2f}%)", False),
        ("Recall Micro (Sensibilidad)", f"{gemma['metricas']['micro']['recall']:.4f} ({gemma['metricas']['micro']['recall']*100:.2f}%)", f"{g35['metricas']['micro']['recall']:.4f} ({g35['metricas']['micro']['recall']*100:.2f}%)", f"{g36['metricas']['micro']['recall']:.4f} ({g36['metricas']['micro']['recall']*100:.2f}%)", False),
        ("Micro-F1-Score [IC 95% Bootstrap]", f"{gemma['metricas']['micro']['f1']:.4f} [{gemma['ci_95']['micro_f1'][0]:.3f}, {gemma['ci_95']['micro_f1'][1]:.3f}]", f"{g35['metricas']['micro']['f1']:.4f} [{g35['ci_95']['micro_f1'][0]:.3f}, {g35['ci_95']['micro_f1'][1]:.3f}]", f"{g36['metricas']['micro']['f1']:.4f} [{g36['ci_95']['micro_f1'][0]:.3f}, {g36['ci_95']['micro_f1'][1]:.3f}]", False),
        
        ("BALANCE ANTE CÓDIGOS MINORITARIOS (NIVEL MACRO)", "", "", "", True),
        ("Precisión Macro", f"{gemma['metricas']['macro']['precision']:.4f}", f"{g35['metricas']['macro']['precision']:.4f}", f"{g36['metricas']['macro']['precision']:.4f}", False),
        ("Recall Macro", f"{gemma['metricas']['macro']['recall']:.4f}", f"{g35['metricas']['macro']['recall']:.4f}", f"{g36['metricas']['macro']['recall']:.4f}", False),
        ("Macro-F1-Score [IC 95% Bootstrap]", f"{gemma['metricas']['macro']['f1']:.4f} [{gemma['ci_95']['macro_f1'][0]:.3f}, {gemma['ci_95']['macro_f1'][1]:.3f}]", f"{g35['metricas']['macro']['f1']:.4f} [{g35['ci_95']['macro_f1'][0]:.3f}, {g35['ci_95']['macro_f1'][1]:.3f}]", f"{g36['metricas']['macro']['f1']:.4f} [{g36['ci_95']['macro_f1'][0]:.3f}, {g36['ci_95']['macro_f1'][1]:.3f}]", False),
        
        ("PONDERACIÓN POR PREVALENCIA REAL (NIVEL WEIGHTED)", "", "", "", True),
        ("Precisión Weighted", f"{gemma['metricas']['weighted']['precision']:.4f}", f"{g35['metricas']['weighted']['precision']:.4f}", f"{g36['metricas']['weighted']['precision']:.4f}", False),
        ("Recall Weighted", f"{gemma['metricas']['weighted']['recall']:.4f}", f"{g35['metricas']['weighted']['recall']:.4f}", f"{g36['metricas']['weighted']['recall']:.4f}", False),
        ("Weighted-F1-Score [IC 95% Bootstrap]", f"{gemma['metricas']['weighted']['f1']:.4f} [{gemma['ci_95']['weighted_f1'][0]:.3f}, {gemma['ci_95']['weighted_f1'][1]:.3f}]", f"{g35['metricas']['weighted']['f1']:.4f} [{g35['ci_95']['weighted_f1'][0]:.3f}, {g35['ci_95']['weighted_f1'][1]:.3f}]", f"{g36['metricas']['weighted']['f1']:.4f} [{g36['ci_95']['weighted_f1'][0]:.3f}, {g36['ci_95']['weighted_f1'][1]:.3f}]", False),
    ]
    return filas


def exportar_documento_docx(resultados: List[Dict[str, Any]], ruta_salida: Path):
    if not TIENE_DOCX:
        print(" [AVISO] python-docx no está disponible. No se pudo generar el archivo .docx.")
        return

    def set_cell_shading(cell, color_hex):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="8" w:space="0" w:color="333333"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="333333"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
        '''
        tblPr.append(parse_xml(borders_xml))

    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Desempeño Diagnóstico en la Codificación CIF Automatizada")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.name = 'Times New Roman'

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Evaluación comparativa de modelos LLM bajo criterio de consenso estricto (3/3) frente al Ground Truth (114 historias clínicas, Core Set CIF de dolor crónico con 27 categorías).")
    run_sub.italic = True
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(6)
    r_h1 = h1.add_run("Tabla 1. Desempeño diagnóstico global y validez multietiqueta por modelo LLM")
    r_h1.bold = True
    r_h1.font.size = Pt(11.5)

    filas_transpuestas = construir_filas_tabla_transpuesta(resultados)
    
    cols1 = ["Variable / Métrica Clínica", "Gemma-4-31B-it", "Gemini Flash 3.5", "Gemini Flash 3.6"]
    widths1 = [Inches(2.6), Inches(1.45), Inches(1.45), Inches(1.45)]

    table1 = doc.add_table(rows=len(filas_transpuestas) + 1, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table1)

    hdr_cells1 = table1.rows[0].cells
    for i, name in enumerate(cols1):
        hdr_cells1[i].text = name
        hdr_cells1[i].width = widths1[i]
        set_cell_shading(hdr_cells1[i], "EFEFEF")
        set_cell_margins(hdr_cells1[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells1[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = 'Times New Roman'

    for row_idx, (metrica, v1, v2, v3, es_seccion) in enumerate(filas_transpuestas, start=1):
        cells = table1.rows[row_idx].cells
        
        if es_seccion:
            cells[0].text = metrica
            cells[1].text = ""
            cells[2].text = ""
            cells[3].text = ""
            for c_i, c in enumerate(cells):
                c.width = widths1[c_i]
                set_cell_shading(c, "F7F9FA")
                set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9.0)
                    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                    r.font.name = 'Times New Roman'
        else:
            valores = [metrica, v1, v2, v3]
            for col_idx, val in enumerate(valores):
                cells[col_idx].text = val
                cells[col_idx].width = widths1[col_idx]
                set_cell_margins(cells[col_idx], top=70, bottom=70, left=100, right=100)
                p = cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.0)
                    r.font.name = 'Times New Roman'
                    if "F1-Score" in metrica or "Exact Match" in metrica:
                        if col_idx == 0:
                            r.bold = True
                        elif col_idx > 0:
                            r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    r_h2 = h2.add_run("Tabla 2. Auditoría detallada per class en las 27 categorías CIF del Core Set de Dolor Crónico")
    r_h2.bold = True
    r_h2.font.size = Pt(11.5)

    gemma_m = next(r["metricas"]["por_clase"] for r in resultados if r["meta"]["id"] == "gemma_31b")
    g36_m = next(r["metricas"]["por_clase"] for r in resultados if r["meta"]["id"] == "gemini_flash_36")
    g35_m = next(r["metricas"]["por_clase"] for r in resultados if r["meta"]["id"] == "gemini_flash_35")

    cols2 = ["Código CIF", "Categoría CIF (Core Set)", "Soporte (GT)", "Gemma Prec.", "Gemma Rec.", "Gemma F1", "Flash 3.6 Prec.", "Flash 3.6 Rec.", "Flash 3.6 F1", "Flash 3.5 F1"]
    widths2 = [Inches(0.55), Inches(2.15), Inches(0.45), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.55)]

    table2 = doc.add_table(rows=len(gemma_m) + 1, cols=len(cols2))
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table2)

    hdr_cells2 = table2.rows[0].cells
    for i, name in enumerate(cols2):
        hdr_cells2[i].text = name
        hdr_cells2[i].width = widths2[i]
        set_cell_shading(hdr_cells2[i], "EFEFEF")
        set_cell_margins(hdr_cells2[i], top=80, bottom=80, left=60, right=60)
        p = hdr_cells2[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in [0, 1] else WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.name = 'Times New Roman'

    for row_idx, c in enumerate(sorted(gemma_m.keys()), start=1):
        cells = table2.rows[row_idx].cells
        nom = gemma_m[c]["nombre"]
        sup = str(gemma_m[c]["soporte"])
        vals = [
            c,
            nom,
            sup,
            f"{gemma_m[c]['precision']:.4f}",
            f"{gemma_m[c]['recall']:.4f}",
            f"{gemma_m[c]['f1']:.4f}",
            f"{g36_m[c]['precision']:.4f}",
            f"{g36_m[c]['recall']:.4f}",
            f"{g36_m[c]['f1']:.4f}",
            f"{g35_m[c]['f1']:.4f}"
        ]
        for col_idx, val in enumerate(vals):
            cells[col_idx].text = val
            cells[col_idx].width = widths2[col_idx]
            set_cell_margins(cells[col_idx], top=60, bottom=60, left=60, right=60)
            p = cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in [0, 1] else WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8.0)
                r.font.name = 'Times New Roman'
                if col_idx == 0:
                    r.bold = True

    doc.save(ruta_salida)
    print(f"📄 Documento Word (.docx) exportado exitosamente a: {ruta_salida}")


def main():
    print("=" * 115)
    print(" 🔬 EVALUACIÓN METODOLÓGICA DE DESEMPEÑO EN CODIFICACIÓN CIF (MICRO, MACRO, WEIGHTED, PER CLASS)")
    print("    Criterio: Consenso Estricto (3/3) frente al Ground Truth (114 Historias | 27 Categorías CIF)")
    print("=" * 115)

    codigos_cif = obtener_todos_los_codigos(MODELOS)
    resultados = []

    for m in MODELOS:
        if not m["archivo"].exists():
            print(f" [ERROR] Archivo no encontrado: {m['archivo']}")
            continue

        with open(m["archivo"], "r", encoding="utf-8") as f:
            historias = json.load(f)

        metricas = calcular_desempeno(historias, codigos_cif)
        ci_95 = calcular_bootstrap_ci(historias, codigos_cif, n_iter=1000, semilla=2026)

        resultados.append({
            "meta": m,
            "metricas": metricas,
            "ci_95": ci_95
        })

    filas = construir_filas_tabla_transpuesta(resultados)
    print(f"\n {'Métrica / Variable Clínica':<46} | {'Gemma-4-31B-it':<20} | {'Gemini Flash 3.5':<20} | {'Gemini Flash 3.6':<20}")
    print("-" * 115)

    for met, v1, v2, v3, es_sec in filas:
        if es_sec:
            print(f"\n▶ {met}")
            print("-" * 115)
        else:
            print(f" {met:<45} | {v1:<20} | {v2:<20} | {v3:<20}")

    print("=" * 115)

    ruta_docx = TABLAS_DIR / "tablas_desempeno.docx"
    exportar_documento_docx(resultados, ruta_docx)

    ruta_json = LLM_DIR / "resumen_f1_score.json"
    salida_json = []
    for r in resultados:
        salida_json.append({
            "modelo_id": r["meta"]["id"],
            "modelo_nombre": r["meta"]["nombre"],
            "metricas": r["metricas"],
            "ci_95": r["ci_95"]
        })
    with open(ruta_json, "w", encoding="utf-8") as f:
        json.dump(salida_json, f, indent=2, ensure_ascii=False)
    print(f"💾 Resumen estadístico JSON guardado en: {ruta_json}")


if __name__ == "__main__":
    main()
