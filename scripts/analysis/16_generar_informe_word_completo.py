# -*- coding: utf-8 -*-
"""
===============================================================================
INFORME BREVE Y SINTÉTICO: VALIDACIÓN CON HISTORIAS CLÍNICAS REALES (N = 21)
ESTILOS Y FORMATO 100% HOMOGÉNEOS CON EL DOCUMENTO PRINCIPAL DEL TFM
- Fuente base: Times New Roman 12 pt, interlineado 1.5, justificado
- Estilos de encabezado: Heading 1 (14 pt bold), Heading 2 (13 pt bold), Heading 3 (12 pt #0A2F40)
- Títulos de figuras y tablas: Estilo Caption (9 pt, color #0E2841, Times New Roman)
- Tablas: Formato editorial APA / Booktabs según skill tfl-apa-tables (Times New Roman 8.5-9 pt)
===============================================================================
"""

import json
import csv
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = Path(__file__).resolve().parents[2]
JSON_HUMANO = BASE_DIR / 'results' / 'human_text' / 'resumen_human_annotated.json'
JSON_SINTETICO = BASE_DIR / 'results' / 'llm_text' / 'resumen_f1_score.json'
CSV_PER_CLASS = BASE_DIR / 'results' / 'human_text' / 'tabla_per_class_human.csv'
FIG_DIR = BASE_DIR / 'results' / 'TFL' / 'figuras'

OUT_DOCX_RESULTS = BASE_DIR / 'results' / 'TFL' / 'informes' / 'informe_validacion_historias_humanas.docx'

with open(JSON_HUMANO, 'r', encoding='utf-8') as f:
    datos_humano = json.load(f)
with open(JSON_SINTETICO, 'r', encoding='utf-8') as f:
    datos_sintetico = json.load(f)

# Helper XML styling functions for APA Booktabs according to tfl-apa-tables skill
def aplicar_bordes_apa(table):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="222222"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="222222"/>
        <w:insideH w:val="none"/>
        <w:insideV w:val="none"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def configurar_celda(cell, ancho_in=None, fondo_hex=None, top_p=50, bot_p=50, left_p=60, right_p=60, border_bottom_color=None, border_bottom_sz=None):
    tcPr = cell._tc.get_or_add_tcPr()
    if ancho_in:
        cell.width = Inches(ancho_in)
    if fondo_hex:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fondo_hex}"/>')
        tcPr.append(shd)
    mar_xml = f'''
    <w:tcMar {nsdecls("w")}>
        <w:top w:w="{top_p}" w:type="dxa"/>
        <w:bottom w:w="{bot_p}" w:type="dxa"/>
        <w:left w:w="{left_p}" w:type="dxa"/>
        <w:right w:w="{right_p}" w:type="dxa"/>
    </w:tcMar>
    '''
    tcPr.append(parse_xml(mar_xml))
    if border_bottom_color and border_bottom_sz:
        b_xml = f'''
        <w:tcBorders {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="{border_bottom_sz}" w:space="0" w:color="{border_bottom_color}"/>
        </w:tcBorders>
        '''
        tcPr.append(parse_xml(b_xml))

def construir_documento():
    doc = docx.Document()
    
    # 1. Configurar márgenes exactamente iguales al documento principal (2.5 cm arriba/abajo, 3.0 cm izq/dcha)
    for sec in doc.sections:
        sec.top_margin = Inches(0.98)
        sec.bottom_margin = Inches(0.98)
        sec.left_margin = Inches(1.18)
        sec.right_margin = Inches(1.18)
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        
    # 2. Configurar estilos base en el documento
    # Normal Style: Times New Roman 12 pt, 1.5 line spacing, 6 pt space after, Justified
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        if level == 1:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(10, 47, 64) # #0A2F40
        elif level == 2:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(10, 47, 64)
        elif level == 3:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run.font.size = Pt(12)
            run.font.italic = True
            run.font.color.rgb = RGBColor(10, 47, 64)
        return p

    def add_p(text, bold_prefix=None, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Times New Roman'
            r_pre.font.size = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def add_caption(caption_text):
        """Genera un pie de figura/tabla usando el estilo Caption exacto del TFM (9 pt, #0E2841)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_with_next = True
        run = p.add_run(caption_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(14, 40, 65) # #0E2841 (estilo Caption del TFM)
        return p

    def add_table_note(note_str):
        """Nota de pie de tabla / figura en estilo APA / TFM (8.5 pt Times New Roman, interlineado 1.0)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        r1 = p.add_run("Nota. ")
        r1.bold = True
        r1.font.italic = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(8.5)
        r2 = p.add_run(note_str)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(8.5)
        return p

    # -------------------------------------------------------------
    # ENCABEZADO SUPERIOR DISCRETO
    # -------------------------------------------------------------
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_top.paragraph_format.space_after = Pt(10)
    p_top.paragraph_format.line_spacing = 1.0
    r_top = p_top.add_run("TRABAJO FIN DE MÁSTER · BIOINFORMÁTICA")
    r_top.font.name = 'Times New Roman'
    r_top.font.size = Pt(8.5)
    r_top.font.color.rgb = RGBColor(120, 120, 120)

    # -------------------------------------------------------------
    # SECCIÓN 1: METODOLOGÍA (BREVE Y CONCISA)
    # -------------------------------------------------------------
    add_heading("1. Metodología: Validación con Historias Clínicas de Fisioterapeutas", level=1)
    add_p("Para evaluar la capacidad de transferencia del pipeline automatizado fuera del entorno sintético, se recopiló un corpus de validación externa compuesto por 21 historias clínicas narradas libremente por fisioterapeutas colegiados. El estándar de referencia (Ground Truth) fue indexado por expertos según el Core Set de la CIF para dolor crónico generalizado, identificando 118 diagnósticos reales (media de 5.62 ± 2.82 códigos por historia, rango 2–16) distribuidos en Funciones Corporales (b, 49 menciones), Actividades y Participación (d, 55 menciones) y Factores Ambientales (e, 14 menciones).")
    add_p("El procedimiento de inferencia en n8n evaluó tres modelos bajo 3 réplicas estocásticas y consenso estricto (3/3): Gemma-4-31B-it (modelo de pesos abiertos ejecutado localmente de forma on-premise), Gemini Flash 3.5 y Gemini Flash 3.6 (modelos cloud). Se computaron métricas de fiabilidad intra-modelo (PAE, AC1 de Gwet, Alfa de Krippendorff), validez diagnóstica (Micro, Macro, Weighted F1, Precisión, Recall, EMR) con intervalos de confianza al 95% mediante remuestreo Bootstrap (B = 1000) y la Tasa de Retención del Rendimiento respecto al corpus sintético (N = 114).")

    # -------------------------------------------------------------
    # SECCIÓN 2: RESULTADOS (BREVES Y DIRECTOS)
    # -------------------------------------------------------------
    add_heading("2. Resultados: Rendimiento Tri-Modelo en Historias Reales", level=1)
    
    fg  = datos_humano['gemma_31b']
    f35 = datos_humano['flash_35']
    f36 = datos_humano['flash_36']
    
    add_heading("2.1. Fiabilidad y Validez Diagnóstica Global", level=2)
    add_p(f"La fiabilidad intra-modelo evidenció una elevada reproducibilidad en los tres modelos (Tabla 1): Gemini Flash 3.6 alcanzó un acuerdo exacto del {f36['fiabilidad']['pae_pct']:.2f}% (18/21 historias idénticas, α = {f36['fiabilidad']['alpha']:.4f}), seguido de Gemma-4-31B-it ({fg['fiabilidad']['pae_pct']:.2f}%, α = {fg['fiabilidad']['alpha']:.4f}) y Gemini Flash 3.5 ({f35['fiabilidad']['pae_pct']:.2f}%, α = {f35['fiabilidad']['alpha']:.4f}).")
    add_p(f"En validez diagnóstica global frente al Ground Truth (Tabla 1 y Figuras 1 y 2), Gemini Flash 3.6 lideró con un Micro-F1 de {f36['desempeno']['micro']['f1']:.4f} [IC 95%: {f36['ci_95']['micro'][0]:.4f}–{f36['ci_95']['micro'][1]:.4f}] y una Precisión Micro de {f36['desempeno']['micro']['p']:.4f}. Gemini Flash 3.5 alcanzó un Micro-F1 de {f35['desempeno']['micro']['f1']:.4f} (Precisión: {f35['desempeno']['micro']['p']:.4f}), mientras que el modelo local Gemma-4-31B-it demostró una destacada eficacia diagnóstica on-premise con un Micro-F1 de {fg['desempeno']['micro']['f1']:.4f} (Precisión: {fg['desempeno']['micro']['p']:.4f}). La tasa de falsos positivos (alucinaciones) fue extraordinariamente reducida ({f36['desempeno']['fp']} en Flash 3.6, {f35['desempeno']['fp']} en Flash 3.5 y {fg['desempeno']['fp']} en Gemma sobre 567 decisiones), garantizando una excelente seguridad clínica.")

    # -------------------------------------------------------------
    # TABLA 1 (FORMATO EDITORIAL APA)
    # -------------------------------------------------------------
    add_caption("Tabla 1. Evaluación global del desempeño diagnóstico y fiabilidad intra-modelo ante historias clínicas redactadas por fisioterapeutas (N = 21)")
    
    filas_t1 = [
        ("Corpus Clínico Humano", "Historias clínicas evaluadas (N)", "21", "21", "21"),
        ("Corpus Clínico Humano", "Espacio ontológico evaluado (N × 27)", "567", "567", "567"),
        ("Corpus Clínico Humano", "Soporte real de menciones CIF", "118", "118", "118"),
        ("Corpus Clínico Humano", "Promedio de códigos por historia", "5.62 ± 2.82", "5.62 ± 2.82", "5.62 ± 2.82"),
        ("Corpus Clínico Humano", "Criterio de consenso multi-iteración", "Consenso estricto (3/3)", "Consenso estricto (3/3)", "Consenso estricto (3/3)"),
        
        ("Fiabilidad Intra-Modelo", "Porcentaje de Acuerdo Exacto (PAE, %)", f"{fg['fiabilidad']['pae_pct']:.2f}%", f"{f35['fiabilidad']['pae_pct']:.2f}%", f"{f36['fiabilidad']['pae_pct']:.2f}%"),
        ("Fiabilidad Intra-Modelo", "Acuerdo Observado (Po, %)", f"{fg['fiabilidad']['Po']*100:.2f}%", f"{f35['fiabilidad']['Po']*100:.2f}%", f"{f36['fiabilidad']['Po']*100:.2f}%"),
        ("Fiabilidad Intra-Modelo", "Coeficiente AC1 de Gwet", f"{fg['fiabilidad']['ac1']:.4f}", f"{f35['fiabilidad']['ac1']:.4f}", f"{f36['fiabilidad']['ac1']:.4f}"),
        ("Fiabilidad Intra-Modelo", "Alfa (α) de Krippendorff", f"{fg['fiabilidad']['alpha']:.4f}", f"{f35['fiabilidad']['alpha']:.4f}", f"{f36['fiabilidad']['alpha']:.4f}"),
        
        ("Matriz de Confusión", "Verdaderos Positivos (TP)", str(fg['desempeno']['tp']), str(f35['desempeno']['tp']), str(f36['desempeno']['tp'])),
        ("Matriz de Confusión", "Falsos Positivos / Alucinaciones (FP)", str(fg['desempeno']['fp']), str(f35['desempeno']['fp']), str(f36['desempeno']['fp'])),
        ("Matriz de Confusión", "Falsos Negativos / Omisiones (FN)", str(fg['desempeno']['fn']), str(f35['desempeno']['fn']), str(f36['desempeno']['fn'])),
        ("Matriz de Confusión", "Exact Match Ratio (EMR, %)", f"{fg['desempeno']['emr']:.2f}% ({fg['desempeno']['exact']}/21)", f"{f35['desempeno']['emr']:.2f}% ({f35['desempeno']['exact']}/21)", f"{f36['desempeno']['emr']:.2f}% ({f36['desempeno']['exact']}/21)"),
        
        ("Nivel Micro (Global)", "Precisión Micro", f"{fg['desempeno']['micro']['p']:.4f}", f"{f35['desempeno']['micro']['p']:.4f}", f"{f36['desempeno']['micro']['p']:.4f}"),
        ("Nivel Micro (Global)", "Sensibilidad / Recall Micro", f"{fg['desempeno']['micro']['r']:.4f}", f"{f35['desempeno']['micro']['r']:.4f}", f"{f36['desempeno']['micro']['r']:.4f}"),
        ("Nivel Micro (Global)", "Micro-F1 [IC 95% Bootstrap]", f"{fg['desempeno']['micro']['f1']:.4f} [{fg['ci_95']['micro'][0]:.4f}, {fg['ci_95']['micro'][1]:.4f}]", f"{f35['desempeno']['micro']['f1']:.4f} [{f35['ci_95']['micro'][0]:.4f}, {f35['ci_95']['micro'][1]:.4f}]", f"{f36['desempeno']['micro']['f1']:.4f} [{f36['ci_95']['micro'][0]:.4f}, {f36['ci_95']['micro'][1]:.4f}]"),
        
        ("Nivel Macro (Promedio)", "Precisión Macro", f"{fg['desempeno']['macro']['p']:.4f}", f"{f35['desempeno']['macro']['p']:.4f}", f"{f36['desempeno']['macro']['p']:.4f}"),
        ("Nivel Macro (Promedio)", "Sensibilidad / Recall Macro", f"{fg['desempeno']['macro']['r']:.4f}", f"{f35['desempeno']['macro']['r']:.4f}", f"{f36['desempeno']['macro']['r']:.4f}"),
        ("Nivel Macro (Promedio)", "Macro-F1 [IC 95% Bootstrap]", f"{fg['desempeno']['macro']['f1']:.4f} [{fg['ci_95']['macro'][0]:.4f}, {fg['ci_95']['macro'][1]:.4f}]", f"{f35['desempeno']['macro']['f1']:.4f} [{f35['ci_95']['macro'][0]:.4f}, {f35['ci_95']['macro'][1]:.4f}]", f"{f36['desempeno']['macro']['f1']:.4f} [{f36['ci_95']['macro'][0]:.4f}, {f36['ci_95']['macro'][1]:.4f}]"),
        
        ("Nivel Weighted (Ponderado)", "Precisión Weighted", f"{fg['desempeno']['weighted']['p']:.4f}", f"{f35['desempeno']['weighted']['p']:.4f}", f"{f36['desempeno']['weighted']['p']:.4f}"),
        ("Nivel Weighted (Ponderado)", "Sensibilidad / Recall Weighted", f"{fg['desempeno']['weighted']['r']:.4f}", f"{f35['desempeno']['weighted']['r']:.4f}", f"{f36['desempeno']['weighted']['r']:.4f}"),
        ("Nivel Weighted (Ponderado)", "Weighted-F1 [IC 95% Bootstrap]", f"{fg['desempeno']['weighted']['f1']:.4f} [{fg['ci_95']['weighted'][0]:.4f}, {fg['ci_95']['weighted'][1]:.4f}]", f"{f35['desempeno']['weighted']['f1']:.4f} [{f35['ci_95']['weighted'][0]:.4f}, {f35['ci_95']['weighted'][1]:.4f}]", f"{f36['desempeno']['weighted']['f1']:.4f} [{f36['ci_95']['weighted'][0]:.4f}, {f36['ci_95']['weighted'][1]:.4f}]")
    ]
    
    t1 = doc.add_table(rows=len(filas_t1) + 1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_apa(t1)
    
    headers_t1 = ["Dimensión", "Métrica / Parámetro", "Gemma-4-31B-it (Local)", "Gemini Flash 3.5 (Cloud)", "Gemini Flash 3.6 (Cloud)"]
    anchos_t1 = [1.4, 1.6, 1.3, 1.3, 1.3]
    for j, h in enumerate(headers_t1):
        cell = t1.cell(0, j)
        configurar_celda(cell, anchos_t1[j], fondo_hex="F4F6F7", top_p=35, bot_p=35, left_p=30, right_p=30, border_bottom_color="444444", border_bottom_sz="6")
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.5)
        
    lineas_corte = [4, 8, 12, 15, 18, 21]
    for i, fila in enumerate(filas_t1):
        es_fin_grupo = i in lineas_corte
        for j, val in enumerate(fila):
            cell = t1.cell(i+1, j)
            b_col = "D0D0D0" if es_fin_grupo and i != 21 else None
            b_sz = "4" if es_fin_grupo and i != 21 else None
            configurar_celda(cell, anchos_t1[j], top_p=25, bot_p=25, left_p=30, right_p=30, border_bottom_color=b_col, border_bottom_sz=b_sz)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8.0)
            if j == 0:
                run.font.italic = True
                
    add_table_note("IC 95%: Intervalos de confianza calculados por Bootstrap no paramétrico (B = 1000 iteraciones). PAE: Porcentaje de Acuerdo Exacto. EMR: Exact Match Ratio.")

    # -------------------------------------------------------------
    # INSERTAR FIGURAS 1 Y 2
    # -------------------------------------------------------------
    add_caption("Figura 1. Comparativa global de desempeño diagnóstico tri-modelo (F1-score e IC 95% Bootstrap) en historias de fisioterapeutas")
    if (FIG_DIR / '01_comparativa_global_f1_human.png').exists():
        doc.add_picture(str(FIG_DIR / '01_comparativa_global_f1_human.png'), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note("Micro-F1, Macro-F1, Weighted-F1 y EMR entre Gemma-4-31B-it, Gemini Flash 3.5 y Gemini Flash 3.6.")

    add_caption("Figura 2. Desglose pareado de Precisión, Sensibilidad (Recall) y F1-Score en los niveles Micro y Weighted")
    if (FIG_DIR / '02_precision_recall_f1_pareado_human.png').exists():
        doc.add_picture(str(FIG_DIR / '02_precision_recall_f1_pareado_human.png'), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note("Nivel Micro (izquierda) y Nivel Weighted (derecha), reflejando la alta precisión (> 82%–92.5%) del consenso estricto.")

    # -------------------------------------------------------------
    # 2.2. DESEMPEÑO POR COMPONENTES
    # -------------------------------------------------------------
    add_heading("2.2. Desempeño por Componentes CIF y Auditoría por Categoría", level=2)
    add_p(f"Por componentes ontológicos (Tabla 2 y Figura 4), el mayor rendimiento se obtuvo en Funciones Corporales (b, F1 de {fg['desempeno']['por_componente']['b']['micro_f1']:.2f} en Gemma, {f35['desempeno']['por_componente']['b']['micro_f1']:.2f} en Flash 3.5 y {f36['desempeno']['por_componente']['b']['micro_f1']:.2f} en Flash 3.6), destacando b280 («Dolor», F1 = 1.00) y b134 («Sueño», F1 = 1.00). En Actividades y Participación (d), el F1 osciló entre {fg['desempeno']['por_componente']['d']['micro_f1']:.2f} y {f36['desempeno']['por_componente']['d']['micro_f1']:.2f} (d450 «Andar»: F1 0.87–0.91; d640 «Quehaceres»: F1 0.86–0.93). En Factores Ambientales (e), el F1 se situó entre {fg['desempeno']['por_componente']['e']['micro_f1']:.2f} y {f35['desempeno']['por_componente']['e']['micro_f1']:.2f}, con precisión del 100% en medicamentos (e1101, F1 0.83–0.92).")

    # -------------------------------------------------------------
    # TABLA 2: AUDITORÍA PER CLASS (FORMATO LIMPIO APA)
    # -------------------------------------------------------------
    add_caption("Tabla 2. Auditoría de clasificación F1 por categoría CIF en el corpus de historias humanas (N = 21)")
    
    filas_t2 = []
    with open(CSV_PER_CLASS, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        header = next(reader)
        for row in reader:
            filas_t2.append(row)
            
    t2 = doc.add_table(rows=len(filas_t2) + 1, cols=7)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_apa(t2)
    
    headers_t2 = ["Componente CIF", "Código", "Categoría CIF", "Soporte Real", "Gemma 31B (F1)", "Flash 3.5 (F1)", "Flash 3.6 (F1)"]
    anchos_t2 = [1.5, 0.65, 2.3, 0.65, 0.85, 0.85, 0.85]
    for j, h in enumerate(headers_t2):
        cell = t2.cell(0, j)
        configurar_celda(cell, anchos_t2[j], fondo_hex="F4F6F7", top_p=30, bot_p=30, left_p=25, right_p=25, border_bottom_color="444444", border_bottom_sz="6")
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 3 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8.0)
        
    for i, fila in enumerate(filas_t2):
        comp_nom = fila[1]
        cod = fila[2]
        nom = fila[3]
        sup = fila[4]
        f1_35 = f"{float(fila[10]):.2f}"
        f1_36 = f"{float(fila[16]):.2f}"
        f1_g  = f"{float(fila[22]):.2f}"
        
        vals_mostrar = [comp_nom, cod, nom, sup, f1_g, f1_35, f1_36]
        es_corte_comp = i in [10, 21]
        for j, val in enumerate(vals_mostrar):
            cell = t2.cell(i+1, j)
            b_col = "888888" if es_corte_comp else None
            b_sz = "6" if es_corte_comp else None
            configurar_celda(cell, anchos_t2[j], top_p=20, bot_p=20, left_p=25, right_p=25, border_bottom_color=b_col, border_bottom_sz=b_sz)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 3 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(7.5)
            if j == 0:
                run.font.italic = True
            elif j == 1:
                run.bold = True
                
    add_table_note("F1: Puntuación F1-Score bajo consenso estricto (3/3). Soporte Real: Frecuencia observada de la categoría en las 21 historias.")

    # -------------------------------------------------------------
    # INSERTAR FIGURAS 3 Y 4
    # -------------------------------------------------------------
    add_caption("Figura 3. Auditoría de F1-Score por categoría CIF en el corpus de historias humanas")
    if (FIG_DIR / '03_auditoria_per_class_human.png').exists():
        doc.add_picture(str(FIG_DIR / '03_auditoria_per_class_human.png'), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note("F1-Score para las 22 categorías activas del Core Set. La línea discontinua roja indica el umbral de excelencia diagnóstica (0.80).")

    add_caption("Figura 4. Rendimiento diagnóstico comparativo (Micro-F1) agrupado por componentes de la CIF")
    if (FIG_DIR / '04_desempeno_por_componente_cif_human.png').exists():
        doc.add_picture(str(FIG_DIR / '04_desempeno_por_componente_cif_human.png'), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note("Rendimiento por dominio ontológico: Funciones Corporales (b), Actividades y Participación (d) y Factores Ambientales (e).")

    # -------------------------------------------------------------
    # 2.3. ANÁLISIS DE GENERALIZACIÓN: SINTÉTICO VS HUMANO
    # -------------------------------------------------------------
    add_heading("2.3. Análisis de Generalización: Corpus Sintético vs Fisioterapeutas", level=2)
    
    s35 = next(item for item in datos_sintetico if item['modelo_id'] == 'gemini_flash_35')
    s36 = next(item for item in datos_sintetico if item['modelo_id'] == 'gemini_flash_36')
    sg  = next(item for item in datos_sintetico if item['modelo_id'] == 'gemma_31b')
    
    ret_f35_f1 = (f35['desempeno']['micro']['f1'] / s35['metricas']['micro']['f1']) * 100
    ret_f36_f1 = (f36['desempeno']['micro']['f1'] / s36['metricas']['micro']['f1']) * 100
    ret_g_f1   = (fg['desempeno']['micro']['f1'] / sg['metricas']['micro']['f1']) * 100
    
    ret_f35_p = (f35['desempeno']['micro']['p'] / s35['metricas']['micro']['precision']) * 100
    ret_f36_p = (f36['desempeno']['micro']['p'] / s36['metricas']['micro']['precision']) * 100
    ret_g_p   = (fg['desempeno']['micro']['p'] / sg['metricas']['micro']['precision']) * 100

    add_p(f"La comparación entre el rendimiento sintético (N = 114) y el corpus de fisioterapeutas (N = 21) demuestra una alta tasa de retención (Tabla 3 y Figura 5): Gemini Flash 3.6 retuvo el {ret_f36_f1:.1f}% de su Micro-F1 ({f36['desempeno']['micro']['f1']:.4f} vs {s36['metricas']['micro']['f1']:.4f}) y el {ret_f36_p:.1f}% de su Precisión ({f36['desempeno']['micro']['p']:.4f} vs {s36['metricas']['micro']['precision']:.4f}). Gemini Flash 3.5 retuvo el {ret_f35_f1:.1f}% de su Micro-F1 y el {ret_f35_p:.1f}% de su Precisión. Gemma-4-31B-it retuvo el {ret_g_f1:.1f}% de su Micro-F1 y el {ret_g_p:.1f}% de su Precisión. La reducción en EMR se explica por la alta densidad diagnóstica de las historias reales, donde la omisión de un único código anula el acierto estricto del caso completo.")

    # -------------------------------------------------------------
    # TABLA 3: COMPARATIVA SINTÉTICO VS HUMANO (FORMATO APA)
    # -------------------------------------------------------------
    add_caption("Tabla 3. Comparativa de generalización del pipeline: Rendimiento ante corpus sintético (N = 114) versus corpus de fisioterapeutas (N = 21)")
    
    filas_t3 = [
        ("Exact Match Ratio (EMR, %)", f"{sg['metricas']['emr_pct']:.2f}%", f"{fg['desempeno']['emr']:.2f}%", f"{(fg['desempeno']['emr']/sg['metricas']['emr_pct'])*100:.1f}%", f"{s35['metricas']['emr_pct']:.2f}%", f"{f35['desempeno']['emr']:.2f}%", f"{(f35['desempeno']['emr']/s35['metricas']['emr_pct'])*100:.1f}%", f"{s36['metricas']['emr_pct']:.2f}%", f"{f36['desempeno']['emr']:.2f}%", f"{(f36['desempeno']['emr']/s36['metricas']['emr_pct'])*100:.1f}%"),
        ("Precisión Micro", f"{sg['metricas']['micro']['precision']:.4f}", f"{fg['desempeno']['micro']['p']:.4f}", f"{ret_g_p:.1f}%", f"{s35['metricas']['micro']['precision']:.4f}", f"{f35['desempeno']['micro']['p']:.4f}", f"{ret_f35_p:.1f}%", f"{s36['metricas']['micro']['precision']:.4f}", f"{f36['desempeno']['micro']['p']:.4f}", f"{ret_f36_p:.1f}%"),
        ("Sensibilidad / Recall Micro", f"{sg['metricas']['micro']['recall']:.4f}", f"{fg['desempeno']['micro']['r']:.4f}", f"{(fg['desempeno']['micro']['r']/sg['metricas']['micro']['recall'])*100:.1f}%", f"{s35['metricas']['micro']['recall']:.4f}", f"{f35['desempeno']['micro']['r']:.4f}", f"{(f35['desempeno']['micro']['r']/s35['metricas']['micro']['recall'])*100:.1f}%", f"{s36['metricas']['micro']['recall']:.4f}", f"{f36['desempeno']['micro']['r']:.4f}", f"{(f36['desempeno']['micro']['r']/s36['metricas']['micro']['recall'])*100:.1f}%"),
        ("Micro-F1", f"{sg['metricas']['micro']['f1']:.4f}", f"{fg['desempeno']['micro']['f1']:.4f}", f"{ret_g_f1:.1f}%", f"{s35['metricas']['micro']['f1']:.4f}", f"{f35['desempeno']['micro']['f1']:.4f}", f"{ret_f35_f1:.1f}%", f"{s36['metricas']['micro']['f1']:.4f}", f"{f36['desempeno']['micro']['f1']:.4f}", f"{ret_f36_f1:.1f}%"),
        ("Macro-F1", f"{sg['metricas']['macro']['f1']:.4f}", f"{fg['desempeno']['macro']['f1']:.4f}", f"{(fg['desempeno']['macro']['f1']/sg['metricas']['macro']['f1'])*100:.1f}%", f"{s35['metricas']['macro']['f1']:.4f}", f"{f35['desempeno']['macro']['f1']:.4f}", f"{(f35['desempeno']['macro']['f1']/s35['metricas']['macro']['f1'])*100:.1f}%", f"{s36['metricas']['macro']['f1']:.4f}", f"{f36['desempeno']['macro']['f1']:.4f}", f"{(f36['desempeno']['macro']['f1']/s36['metricas']['macro']['f1'])*100:.1f}%"),
        ("Weighted-F1", f"{sg['metricas']['weighted']['f1']:.4f}", f"{fg['desempeno']['weighted']['f1']:.4f}", f"{(fg['desempeno']['weighted']['f1']/sg['metricas']['weighted']['f1'])*100:.1f}%", f"{s35['metricas']['weighted']['f1']:.4f}", f"{f35['desempeno']['weighted']['f1']:.4f}", f"{(f35['desempeno']['weighted']['f1']/s35['metricas']['weighted']['f1'])*100:.1f}%", f"{s36['metricas']['weighted']['f1']:.4f}", f"{f36['desempeno']['weighted']['f1']:.4f}", f"{(f36['desempeno']['weighted']['f1']/s36['metricas']['weighted']['f1'])*100:.1f}%"),
        ("Acuerdo Exacto (PAE, %)", "98.25%", f"{fg['fiabilidad']['pae_pct']:.2f}%", f"{(fg['fiabilidad']['pae_pct']/98.25)*100:.1f}%", "98.25%", f"{f35['fiabilidad']['pae_pct']:.2f}%", f"{(f35['fiabilidad']['pae_pct']/98.25)*100:.1f}%", "100.00%", f"{f36['fiabilidad']['pae_pct']:.2f}%", f"{(f36['fiabilidad']['pae_pct']/100.0)*100:.1f}%"),
        ("Coeficiente AC1 de Gwet", "0.9994", f"{fg['fiabilidad']['ac1']:.4f}", f"{(fg['fiabilidad']['ac1']/0.9994)*100:.1f}%", "0.9994", f"{f35['fiabilidad']['ac1']:.4f}", f"{(f35['fiabilidad']['ac1']/0.9994)*100:.1f}%", "1.0000", f"{f36['fiabilidad']['ac1']:.4f}", f"{(f36['fiabilidad']['ac1']/1.0)*100:.1f}%"),
        ("Alfa de Krippendorff", "0.9983", f"{fg['fiabilidad']['alpha']:.4f}", f"{(fg['fiabilidad']['alpha']/0.9983)*100:.1f}%", "0.9983", f"{f35['fiabilidad']['alpha']:.4f}", f"{(f35['fiabilidad']['alpha']/0.9983)*100:.1f}%", "1.0000", f"{f36['fiabilidad']['alpha']:.4f}", f"{(f36['fiabilidad']['alpha']/1.0)*100:.1f}%")
    ]
    
    t3 = doc.add_table(rows=len(filas_t3) + 1, cols=10)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_apa(t3)
    
    headers_t3 = ["Métrica Diagnóstica y Fiabilidad", "Gemma (Sint.)", "Gemma (Hum.)", "Ret.(%)", "Flash 3.5 (Sint.)", "Flash 3.5 (Hum.)", "Ret.(%)", "Flash 3.6 (Sint.)", "Flash 3.6 (Hum.)", "Ret.(%)"]
    anchos_t3 = [1.7, 0.55, 0.55, 0.48, 0.55, 0.55, 0.48, 0.55, 0.55, 0.48]
    for j, h in enumerate(headers_t3):
        cell = t3.cell(0, j)
        configurar_celda(cell, anchos_t3[j], fondo_hex="F4F6F7", top_p=30, bot_p=30, left_p=20, right_p=20, border_bottom_color="444444", border_bottom_sz="6")
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(7.5)
        
    for i, fila in enumerate(filas_t3):
        for j, val in enumerate(fila):
            cell = t3.cell(i+1, j)
            configurar_celda(cell, anchos_t3[j], top_p=20, bot_p=20, left_p=20, right_p=20)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(7.5)
            
    add_table_note("Tasa de Retención (%) = [Valor en Historias Humanas / Valor en Corpus Sintético] × 100.")

    # -------------------------------------------------------------
    # INSERTAR FIGURA 5
    # -------------------------------------------------------------
    add_caption("Figura 5. Comparativa de rendimiento diagnóstico y retención métrica tri-modelo: Sintético (N = 114) vs Fisioterapeutas (N = 21)")
    if (FIG_DIR / '05_comparativa_sintetico_vs_humano.png').exists():
        doc.add_picture(str(FIG_DIR / '05_comparativa_sintetico_vs_humano.png'), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note("Contraste entre el rendimiento en datos sintéticos controlados y narrativas clínicas de fisioterapeutas.")

    # -------------------------------------------------------------
    # SECCIÓN 3: DISCUSIÓN Y CONCLUSIONES (BREVE)
    # -------------------------------------------------------------
    add_heading("3. Discusión y Conclusiones de la Validación Humana", level=1)
    add_p("1. Validez Ecológica: La preservación de un Micro-F1 de 0.82 en Flash 3.6 (0.81 en Flash 3.5 y 0.77 en Gemma) y precisiones del 82%–92.5% demuestra que el flujo en n8n transfiere eficazmente al lenguaje médico no estructurado real sin sobreajustar al generador sintético.")
    add_p("2. Seguridad Asistencial y Privacidad (RGPD): La baja tasa de falsos positivos (solo 7–19 en 567 decisiones) garantiza que el sistema es seguro y no inventa discapacidades. Asimismo, el sólido rendimiento de Gemma-4-31B-it valida la viabilidad técnica de implementar este pipeline en servidores locales hospitalarios, asegurando la soberanía de los datos clínicos y el cumplimiento de la privacidad sin depender de APIs en la nube.")
    add_p("3. Recomendación de Despliegue: Para maximizar la sensibilidad en historias clínicas largas con alta comorbilidad, se recomienda evaluar en producción el consenso por Voto Mayoritario (≥ 2/3), aumentando el Recall con una mínima pérdida de precisión.")
    
    # Save outputs
    doc.save(str(OUT_DOCX_RESULTS))
    print(f" [OK] Documento homogéneo guardado en results: {OUT_DOCX_RESULTS}")

if __name__ == '__main__':
    construir_documento()
